#!/usr/bin/env python3
import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "report_assets"


SECTION_5 = [
    "В ходе преддипломной практики был разработан подход к оценке риска сайта, основанный на объединении нескольких независимых групп признаков. Такой вариант выбран потому, что современные фишинговые страницы часто используют HTTPS, действительные TLS-сертификаты, аккуратный визуальный интерфейс и домены, похожие на легитимные. Поэтому простая проверка адреса или поиск отдельных ключевых слов не обеспечивает достаточной точности.",
    "В разработанном расширении Web Safe итоговый риск определяется на основе URL-признаков, репутационных источников, DNS- и RDAP-сигналов, результатов Google Safe Browsing, а также анализа содержимого активной веб-страницы. Каждый источник данных рассматривается как отдельный слой проверки. Если один из источников временно недоступен, остальные продолжают участвовать в расчёте, что повышает устойчивость решения.",
    "URL-анализ используется для выявления признаков, характерных для фишинговых адресов: небезопасный протокол HTTP, IP-адрес вместо домена, punycode, большое количество поддоменов, необычная длина адреса, высокая доля цифр, наличие дефисов и слов, связанных с авторизацией или восстановлением аккаунта. Эти признаки используются как в эвристическом скоринге, так и в обучаемой логистической модели.",
    "Репутационный анализ выполняется с использованием OpenPhish Community Feed и Phishing.Database. Эти источники позволяют обнаруживать уже известные фишинговые URL и домены. Для снижения количества сетевых запросов данные кэшируются локально. Дополнительно предусмотрена интеграция с Google Safe Browsing через backend-сервер, чтобы API-ключ не попадал в клиентский код расширения.",
    "Сетевые признаки позволяют оценивать свойства доменной инфраструктуры. Через DNS over HTTPS проверяется наличие A/AAAA и NS-записей, количество адресов и минимальное значение TTL. Через RDAP определяется возраст домена и дата последнего изменения регистрационных данных. Молодые домены, недавние изменения регистрационных данных и нетипичные DNS-параметры рассматриваются как дополнительные факторы риска.",
    "Анализ содержимого страницы применяется для выявления признаков социальной инженерии и опасного сбора данных. Расширение проверяет наличие форм, полей пароля, чувствительных полей ввода, внешних обработчиков форм, внешних iframe и сторонних скриптов. Также учитывается наличие брендовой подмены: если страница упоминает известный сервис, но домен не относится к официальным доменам этого сервиса, риск повышается.",
    "Итоговый результат представляется пользователю в виде балла от 0 до 100 и текстового уровня риска: низкий, средний или высокий. Важной особенностью является объяснимость результата: пользователь видит не только итоговую оценку, но и список факторов, которые повлияли на расчёт.",
]

SECTION_6 = [
    "Расширение Web Safe построено по модульному принципу. Такой подход позволяет независимо развивать отдельные части системы: пользовательский интерфейс, сбор признаков страницы, обработку сетевых сигналов, проверку внешних баз угроз и расчёт итогового риска.",
    "Файл manifest.json определяет базовые параметры расширения, разрешения, popup-интерфейс и фоновые сценарии. Для кроссбраузерности используется Manifest V3: в Chromium-браузерах применяется service worker, а для Firefox предусмотрены background scripts. Это позволяет тестировать расширение в Mozilla Firefox и сохранять возможность дальнейшего запуска в Chrome и Edge.",
    "Фоновый модуль background.js выполняет роль координатора. Он получает активную вкладку, внедряет небольшой сценарий сбора признаков страницы, запускает проверки по внешним источникам и передаёт результат в riskEngine.js. После анализа он сохраняет историю проверок и устанавливает badge на иконке расширения.",
    "Модуль riskEngine.js отвечает за формирование итоговой оценки. Он получает все собранные признаки, применяет веса и правила, нормализует результат в диапазон от 0 до 100 и формирует список объяснений для пользователя.",
    "Модуль urlFeatureModel.js вычисляет вероятность риска по структуре URL. Веса модели хранятся в urlModelWeights.js и обновляются при обучении. Обучение выполняется скриптом train_url_model.py на основе фишинговых адресов из OpenPhish и Phishing.Database, а также легитимных доменов из рейтинга Tranco.",
    "Модуль domainUtils.js реализует функции нормализации домена, выделения регистрируемого домена и проверки IP-адресов. Для корректной работы с доменными зонами применяется Public Suffix List, который генерируется в publicSuffixData.js. Это важно, поскольку для доменов вида example.co.uk регистрируемым доменом является example.co.uk, а не co.uk.",
    "Модуль threatFeeds.js загружает и кэширует открытые фиды фишинговых URL и доменов. Модуль networkSignals.js выполняет DNS- и RDAP-проверки. Модуль safeBrowsing.js обращается к backend для проверки URL в Google Safe Browsing.",
    "Пользовательский интерфейс реализован в файлах popup.html, popup.css и popup.js. В интерфейсе отображаются итоговый балл, уровень риска, базовые факты о сайте, статус Google Safe Browsing, результаты URL-модели, признаки страницы, причины оценки и история последних проверок.",
]

SECTION_7 = [
    "Алгоритм работы расширения начинается с открытия пользователем popup-интерфейса. Popup автоматически отправляет сообщение фоновому модулю с запросом на анализ активной вкладки. Повторный анализ можно запустить вручную с помощью кнопки в интерфейсе.",
    "На первом этапе background.js получает активную вкладку и проверяет, что её URL относится к обычной веб-странице с протоколом HTTP или HTTPS. Если пользователь открыл внутреннюю страницу браузера, расширение не выполняет анализ и сообщает об ошибке.",
    "На втором этапе расширение собирает DOM-признаки страницы. Для этого анализируются заголовок страницы, часть видимого текста, формы, поля ввода, iframe, ссылки и внешние скрипты. Особое внимание уделяется формам, которые содержат пароль или другие чувствительные поля и отправляют данные на внешний домен.",
    "На третьем этапе выполняется проверка по открытым фидам угроз. URL и домен сравниваются с локально кэшированными списками OpenPhish и Phishing.Database. Совпадение с такими источниками считается сильным признаком опасности.",
    "На четвёртом этапе выполняется запрос к backend Google Safe Browsing. Расширение не содержит API-ключа и обращается только к собственному backend. Backend, в свою очередь, выполняет запрос к Google Safe Browsing API. Если backend не настроен, расширение продолжает работать по остальным признакам.",
    "На пятом этапе собираются DNS- и RDAP-данные. DNS over HTTPS используется для проверки адресных записей и NS-записей, а RDAP — для определения возраста домена и даты последнего изменения регистрационных данных.",
    "На шестом этапе URL передаётся в обучаемую модель. Модель вычисляет вероятность риска на основе признаков структуры адреса. Такой подход позволяет учитывать не один отдельный признак, а их сочетание.",
    "На заключительном этапе riskEngine.js объединяет все сигналы, формирует итоговый балл риска и список причин. Popup отображает результат пользователю, а background.js сохраняет краткую запись в истории проверок.",
]

SAFE_DEV = {
    "3 Безопасная разработка": [
        "Безопасная разработка является важной частью проекта, поскольку браузерное расширение получает доступ к содержимому активной страницы и может обрабатывать потенциально чувствительные данные. Ошибки в таком программном средстве способны привести не только к некорректной оценке риска, но и к нарушению конфиденциальности пользователя.",
    ],
    "3.1 Планирование требований безопасности": [
        "На этапе планирования были определены основные защищаемые активы: URL посещаемого сайта, признаки DOM-структуры страницы, история проверок, API-ключ Google Safe Browsing и данные внешних фидов угроз. Для каждого актива были определены возможные риски и меры снижения.",
        "Ключевым требованием стало минимальное обращение с пользовательскими данными. Расширение не отправляет содержимое страницы во внешние сервисы. Внешним источникам передаётся только URL, необходимый для проверки. Видимый текст страницы используется локально для оценки риска и не сохраняется в истории.",
    ],
    "3.2 Управление секретами и внешними API": [
        "API-ключ Google Safe Browsing не размещается в исходном коде расширения. Для проверки используется backend safe_browsing_proxy.py, который получает ключ из переменной окружения GOOGLE_SAFE_BROWSING_API_KEY. Такой подход позволяет избежать раскрытия ключа при публикации расширения.",
        "В репозитории добавлен файл .gitignore, исключающий .env, dist, __pycache__ и скомпилированные Python-файлы. Это снижает вероятность случайного попадания секретов, временных файлов и сборочных артефактов в систему контроля версий.",
    ],
    "3.3 Безопасное кодирование расширения": [
        "При разработке расширения применялся принцип минимально необходимых разрешений. В manifest.json указаны только те разрешения, которые требуются для анализа активной вкладки, выполнения скрипта сбора признаков, хранения истории и сетевых запросов к источникам данных.",
        "В пользовательском интерфейсе не используется вставка непроверенного HTML через innerHTML для данных, полученных со страницы. Списки причин и элементы истории формируются через createElement и textContent, что снижает риск XSS внутри popup-интерфейса.",
        "Сбор признаков страницы выполняется ограниченным скриптом, который извлекает только агрегированные характеристики: количество форм, наличие полей пароля, домены action-обработчиков, долю внешних ссылок и другие технические показатели. Полные значения введённых пользователем данных не считываются.",
    ],
    "3.4 Управление зависимостями и сборкой": [
        "Основная часть расширения реализована на JavaScript без подключения сторонних npm-зависимостей. Это уменьшает риск уязвимостей в цепочке поставки. Для вспомогательных скриптов используется Python и стандартная библиотека, а внешние пакеты применяются только для работы с документами и изображениями при подготовке отчёта.",
        "Сборочный процесс реализован в build.py. Он обновляет Public Suffix List, переобучает URL-модель, формирует папку dist/web-safe и архив dist/web-safe.zip. Такой подход позволяет получать воспроизводимый набор файлов для тестирования и распространения.",
    ],
    "3.5 Контроль качества и проверка безопасности": [
        "Для контроля корректности используются проверки manifest.json, проверка импортов JavaScript-модулей, компиляция Python-скриптов и smoke-тестирование URL-модели. Эти проверки позволяют обнаруживать ошибки сборки до загрузки расширения в браузер.",
        "Дополнительно проверяется отказоустойчивость backend Google Safe Browsing. При отсутствии API-ключа backend возвращает корректный ответ о неготовности сервиса, а расширение продолжает анализ по локальным и открытым источникам. Это предотвращает полную остановку анализа при отказе одного компонента.",
    ],
}

SECTION_8 = [
    "Тестирование разработанного средства проводилось по нескольким направлениям: проверка корректности сборки, проверка URL-модели, проверка backend Google Safe Browsing, анализ работы popup-интерфейса и функциональная проверка в Mozilla Firefox.",
    "Для проверки сборки использовался скрипт build.py. Он последовательно обновляет Public Suffix List, переобучает URL-модель и формирует готовый архив расширения dist/web-safe.zip. Последняя сборка завершилась успешно и сформировала пакет расширения для временной загрузки в Firefox.",
    "URL-модель обучалась на 8000 фишинговых и 8000 легитимных примерах. Для снижения завышения качества разделение train/test выполнялось по доменам. По результатам последней сборки модель показала accuracy 0,9904, precision 1,0 и recall 0,9814. Количество ложноположительных срабатываний составило 0, количество пропущенных фишинговых URL — 32.",
    "Для smoke-тестирования модели использовались отдельные URL. Адрес https://google.com/ получил низкую оценку риска около 3 %, что соответствует ожидаемому поведению для легитимного сайта. Адрес http://192.168.0.1/login получил оценку 100 %, так как содержит сразу несколько риск-факторов.",
    "Проверка backend Google Safe Browsing выполнялась через endpoint /health и тестовый POST-запрос /safe-browsing/check. При отсутствии API-ключа backend корректно сообщает, что ключ не задан. Это подтверждает, что расширение может работать в деградированном режиме без полной остановки анализа.",
    "Функциональное тестирование расширения выполнялось через временную загрузку manifest.json в Mozilla Firefox на странице about:debugging. Проверялись автоматический запуск анализа при открытии popup, повторная проверка, отображение причин риска, сохранение истории и установка badge на иконке расширения.",
    "Результаты тестирования показывают, что Web Safe способен выявлять известные фишинговые ресурсы и подозрительные страницы, которые обладают характерными признаками риска. Система не ограничивается поиском ключевых слов, а применяет совокупный анализ URL, домена, DNS/RDAP, репутационных источников и DOM-структуры страницы.",
    "Ограничениями текущей версии являются зависимость части проверок от внешних источников, необходимость дальнейшего расширения обучающей выборки и отсутствие production-backend. Эти ограничения определяют направления дальнейшего развития проекта.",
]

CONCLUSION = [
    "В ходе преддипломной практики было разработано программное средство Web Safe, предназначенное для оценки риска веб-сайтов при работе пользователя в браузере. Разработанное расширение объединяет несколько групп признаков: структуру URL, открытые базы известных фишинговых ресурсов, DNS- и RDAP-сигналы, результаты Google Safe Browsing через backend и анализ содержимого веб-страницы.",
    "Была спроектирована модульная архитектура расширения, реализованы механизмы сбора и обработки признаков, разработан пользовательский popup-интерфейс, добавлена история проверок и предусмотрена сборка проекта в готовый архив расширения. Отдельное внимание уделено безопасной разработке: управлению секретами, минимизации собираемых данных, безопасному формированию интерфейса и воспроизводимой сборке.",
    "Проведённое тестирование подтвердило работоспособность предложенного подхода. URL-модель показала высокие показатели качества на тестовой выборке, а расширение корректно формирует объяснимую оценку риска для пользователя. Разработанное средство может использоваться как основа для дальнейшего развития антифишинговой защиты веб-браузера.",
]

SOURCES = [
    "1. Федеральный закон от 27.07.2006 № 149-ФЗ «Об информации, информационных технологиях и о защите информации».",
    "2. Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных».",
    "3. ГОСТ Р 50922–2006. Защита информации. Основные термины и определения.",
    "4. ГОСТ Р 56939–2024. Защита информации. Разработка безопасного программного обеспечения.",
    "5. OWASP Foundation. OWASP Top 10 Web Application Security Risks.",
    "6. Google Developers. Safe Browsing API Documentation.",
    "7. Mozilla Developer Network. WebExtensions API Documentation.",
    "8. Public Suffix List. Public suffix rules and domain parsing data.",
    "9. OpenPhish. Community Phishing Feed.",
    "10. Phishing.Database. Open phishing domain and URL datasets.",
    "11. Tranco. A Research-Oriented Top Sites Ranking Hardened Against Manipulation.",
]


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("output")
    args = parser.parse_args()

    doc = Document(args.source)
    configure_document(doc)
    replace_document_type(doc)

    figure = 2
    table = 1

    cursor = insert_section(doc, "Предлагаемый подход к оценке риска сайта", SECTION_5)
    cursor, table = add_table_after(
        cursor,
        f"Таблица {table} – Группы сигналов, используемые при оценке риска",
        ["Группа сигналов", "Примеры признаков", "Назначение"],
        [
            ["URL-признаки", "HTTPS, длина URL, поддомены, цифры, punycode", "Выявление подозрительной структуры адреса"],
            ["Репутация", "OpenPhish, Phishing.Database, Google Safe Browsing", "Проверка известных угроз"],
            ["DNS/RDAP", "A/AAAA, NS, TTL, возраст домена", "Оценка инфраструктуры домена"],
            ["DOM-анализ", "формы, пароли, iframe, брендовая подмена", "Выявление опасного поведения страницы"],
        ],
        table,
    )
    cursor, figure = add_picture_after(cursor, "page_signals.png", figure, "Схема анализа содержимого веб-страницы")

    cursor = insert_section(doc, "Архитектура расширения", SECTION_6)
    cursor, figure = add_picture_after(cursor, "architecture.png", figure, "Архитектура расширения Web Safe")
    cursor, figure = add_picture_after(cursor, "project_structure.png", figure, "Структура проекта Web Safe")
    cursor, table = add_table_after(
        cursor,
        f"Таблица {table} – Основные модули программного средства",
        ["Модуль", "Назначение"],
        [
            ["background.js", "Координация анализа активной вкладки"],
            ["riskEngine.js", "Расчёт итоговой оценки риска"],
            ["threatFeeds.js", "Работа с открытыми фидами угроз"],
            ["networkSignals.js", "DNS- и RDAP-анализ домена"],
            ["safeBrowsing.js", "Обмен с backend Google Safe Browsing"],
            ["urlFeatureModel.js", "Вычисление риска по структуре URL"],
        ],
        table,
    )

    cursor = insert_section(doc, "Алгоритм работы расширения", SECTION_7)
    cursor, figure = add_picture_after(cursor, "safe_browsing.png", figure, "Схема интеграции Google Safe Browsing через backend")
    cursor, figure = add_picture_after(cursor, "popup.png", figure, "Интерфейс popup расширения после анализа сайта")

    cursor = insert_custom_section_after(cursor, SAFE_DEV)
    cursor, figure = add_picture_after(cursor, "secure_dev.png", figure, "Контур безопасной разработки проекта")
    cursor, figure = add_picture_after(cursor, "build_pipeline.png", figure, "Сборочный pipeline проекта Web Safe")
    cursor, table = add_table_after(
        cursor,
        f"Таблица {table} – Меры безопасной разработки",
        ["Риск", "Мера снижения"],
        [
            ["Раскрытие API-ключа", "Использование backend и переменных окружения"],
            ["XSS в popup", "Вывод данных через textContent и createElement"],
            ["Избыточный сбор данных", "Сохранение только агрегированных признаков"],
            ["Ошибки сборки", "build.py, проверка manifest и импортов"],
            ["Некорректное выделение домена", "Использование Public Suffix List"],
        ],
        table,
    )

    cursor = insert_section(doc, "Тестирование и оценка эффективности разработанного средства", SECTION_8)
    cursor, figure = add_picture_after(cursor, "model_training.png", figure, "Результат обучения URL-модели и сборки расширения")
    cursor, table = add_table_after(
        cursor,
        f"Таблица {table} – Основные результаты проверки",
        ["Проверка", "Результат"],
        [
            ["Manifest", "JSON корректен"],
            ["Импорты JS", "Все локальные модули найдены"],
            ["Python-скрипты", "py_compile без ошибок"],
            ["URL-модель", "accuracy 0,9904; precision 1,0; recall 0,9814"],
            ["Сборка", "dist/web-safe.zip сформирован"],
        ],
        table,
    )

    cursor = append_heading(doc, cursor, "ЗАКЛЮЧЕНИЕ")
    for item in CONCLUSION:
        cursor = append_body(cursor, item)

    cursor = append_heading(doc, cursor, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    for item in SOURCES:
        cursor = append_body(cursor, item, first_line=False)

    doc.save(args.output)
    print("report generated")


def configure_document(doc):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)


def replace_document_type(doc):
    replacements = {
        "БАКАЛАВРСКАЯ РАБОТА": "ОТЧЁТ О ПРЕДДИПЛОМНОЙ ПРАКТИКЕ",
        "междисциплинарного курсового проекта": "преддипломной практики",
        "междисциплинарного курсового проекта": "преддипломной практики",
    }
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                paragraph.text = paragraph.text.replace(old, new)
                format_body(paragraph)


def find_heading(doc, heading_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == heading_text:
            return paragraph
    raise RuntimeError(f"Heading not found: {heading_text}")


def insert_section(doc, heading_text, paragraphs):
    heading = find_heading(doc, heading_text)
    format_heading(heading)
    cursor = heading
    for text in paragraphs:
        cursor = insert_body_after(cursor, text)
    return cursor


def insert_custom_section_after(cursor, sections):
    for heading, paragraphs in sections.items():
        cursor = insert_paragraph_after(cursor, heading)
        format_heading(cursor)
        for text in paragraphs:
            cursor = insert_body_after(cursor, text)
    return cursor


def insert_body_after(cursor, text):
    paragraph = insert_paragraph_after(cursor, text)
    format_body(paragraph)
    return paragraph


def append_heading(doc, cursor, text):
    paragraph = insert_paragraph_after(cursor, text)
    format_heading(paragraph)
    return paragraph


def append_body(cursor, text, first_line=True):
    paragraph = insert_paragraph_after(cursor, text)
    format_body(paragraph, first_line=first_line)
    return paragraph


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    new_paragraph._element = new_p
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def add_picture_after(cursor, image_name, figure_number, caption):
    image_path = ASSET_DIR / image_name
    paragraph = insert_paragraph_after(cursor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    caption_paragraph = insert_paragraph_after(paragraph, f"Рисунок {figure_number} – {caption}")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_runs(caption_paragraph, 14)
    return caption_paragraph, figure_number + 1


def add_table_after(cursor, caption, headers, rows, table_number):
    caption_paragraph = insert_paragraph_after(cursor, caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_runs(caption_paragraph, 14, bold=False)

    doc = cursor.part.document
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        format_cell(cell, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            format_cell(cells[index])

    caption_paragraph._p.addnext(table._tbl)
    spacer = insert_paragraph_after(caption_paragraph, "")
    table._tbl.addnext(spacer._p)
    return spacer, table_number + 1


def format_cell(cell, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = bold


def format_body(paragraph, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if first_line else Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    format_runs(paragraph, 14)


def format_heading(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    format_runs(paragraph, 14, bold=True)


def format_runs(paragraph, size, bold=None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold


if __name__ == "__main__":
    main()
